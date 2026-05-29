from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any


def _get_param(env, key: str, default: Any = "") -> Any:
    """Read Odoo ir.config_parameter first, then environment variables."""
    if env is not None:
        value = env["ir.config_parameter"].sudo().get_param(key)
        if value not in (None, ""):
            return value

    env_key = key.split(".")[-1].upper()
    aliases = {
        "USE_AI_ANALYSIS": "USE_LLM_ANALYSIS",
        "LLM_PROVIDER": "LLM_PROVIDER",
        "OLLAMA_BASE_URL": "OLLAMA_BASE_URL",
        "OLLAMA_MODEL": "OLLAMA_MODEL",
        "OLLAMA_TEMPERATURE": "OLLAMA_TEMPERATURE",
        "OLLAMA_TIMEOUT": "OLLAMA_TIMEOUT",
    }
    env_name = aliases.get(env_key, env_key)
    return os.getenv(env_name, default)


def _get_bool_param(env, key: str, default: bool = False) -> bool:
    value = str(_get_param(env, key, "true" if default else "false")).strip().lower()
    return value in {"1", "true", "yes", "on", "y"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        import fitz

        doc = fitz.open(str(path))
        try:
            return "\n".join(page.get_text("text") for page in doc)
        finally:
            doc.close()

    if suffix == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True, read_only=True)
        try:
            ws = wb.active
            rows = []
            for row in ws.iter_rows(max_row=50, values_only=True):
                values = [str(v) for v in row if v is not None]
                if values:
                    rows.append(" | ".join(values))
            return "\n".join(rows)
        finally:
            wb.close()

    return ""


def heuristic_suggestions(text: str) -> dict[str, Any]:
    lowered = text.lower()
    suggestions: dict[str, Any] = {
        "company": "",
        "recipient_name": "",
        "recipient_company": "",
        "subject": "",
        "document_type": "Internal Memo",
        "department_prefix": "COO",
        "property_scope": "General",
        "properties": "",
    }

    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.lower().startswith(("re:", "subject:")):
            suggestions["subject"] = clean.split(":", 1)[-1].strip()
            break

    if not suggestions["subject"]:
        first_non_empty = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
        suggestions["subject"] = first_non_empty[:120]

    if "dcl" in lowered or "dunhill" in lowered:
        suggestions["company"] = "DCL"

    if any(word in lowered for word in ["legal", "notice", "demand", "termination"]):
        suggestions["department_prefix"] = "LGL"
        suggestions["document_type"] = "Legal Notice" if "notice" in lowered else "Demand Letter"
    elif any(word in lowered for word in ["invoice", "payment", "rent", "statement", "account"]):
        suggestions["department_prefix"] = "ACC"
        suggestions["document_type"] = "Tenant Letter"
    elif any(word in lowered for word in ["hr", "employee", "leave", "disciplinary"]):
        suggestions["department_prefix"] = "HR"
        suggestions["document_type"] = "HR Letter"
    elif any(word in lowered for word in ["supplier", "quotation", "contract award", "lpo"]):
        suggestions["document_type"] = "Supplier Letter"

    if any(word in lowered for word in ["head office", "internal memo", "memo"]):
        suggestions["property_scope"] = "Head Office"
    elif any(word in lowered for word in ["all properties", "all tenants"]):
        suggestions["property_scope"] = "All Properties"

    return suggestions


def _extract_json(raw: str) -> dict[str, Any]:
    raw = (raw or "").strip()
    if raw.startswith("```json"):
        raw = raw[7:].strip()
    if raw.startswith("```"):
        raw = raw[3:].strip()
    if raw.endswith("```"):
        raw = raw[:-3].strip()

    try:
        return json.loads(raw)
    except Exception:
        pass

    try:
        start = raw.index("{")
        end = raw.rindex("}") + 1
        return json.loads(raw[start:end])
    except Exception:
        return {"raw_response": raw}


def _classification_prompt(text: str) -> str:
    return f"""
You are helping classify a formal business document for an Odoo document reference register.
Return ONLY valid JSON. Do not use markdown.

Return these exact keys:
company, recipient_name, recipient_company, subject, document_type, department_prefix, property_scope, properties.

Rules:
- Use short values.
- If unknown, use an empty string.
- company should be a short company code, for example DCL, if visible or obvious.
- properties can be a comma-separated string if multiple properties are mentioned.
- department_prefix should be one of: COO, DIR, LGL, ACC, HR, OPS, PM, FAC, or empty string.
- property_scope should be one of: Single Property, Multiple Properties, All Properties, Head Office, General, Third Party, None.

Document text:
{text[:9000]}
""".strip()


def analyze_with_ollama(text: str, env=None) -> dict[str, Any]:
    import requests

    base_url = str(
        _get_param(
            env,
            "document_reference_management.ollama_base_url",
            "http://127.0.0.1:11434",
        )
    ).strip().rstrip("/")
    model = str(
        _get_param(
            env,
            "document_reference_management.ollama_model",
            "llama3.1",
        )
    ).strip() or "llama3.1"
    temperature = float(
        _get_param(
            env,
            "document_reference_management.ollama_temperature",
            "0.1",
        )
    )
    timeout = int(
        _get_param(
            env,
            "document_reference_management.ollama_timeout",
            "60",
        )
    )

    prompt = _classification_prompt(text)

    chat_payload = {
        "model": model,
        "stream": False,
        "messages": [
            {
                "role": "system",
                "content": "You extract structured metadata from business documents and return only valid JSON.",
            },
            {"role": "user", "content": prompt},
        ],
        "options": {"temperature": temperature},
    }

    response = requests.post(f"{base_url}/api/chat", json=chat_payload, timeout=timeout)

    # Some older/proxied Ollama setups may not expose /api/chat. Fall back to /api/generate.
    if response.status_code == 404:
        generate_payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": temperature},
        }
        response = requests.post(f"{base_url}/api/generate", json=generate_payload, timeout=timeout)

    response.raise_for_status()
    data = response.json()

    raw = ""
    if isinstance(data, dict):
        if isinstance(data.get("message"), dict):
            raw = data["message"].get("content") or ""
        raw = raw or data.get("response") or ""

    return _extract_json(raw)


def analyze_document(path: Path, env=None) -> dict[str, Any]:
    text = extract_text(path)
    result = heuristic_suggestions(text)
    result["extracted_text_preview"] = text[:1000]

    use_llm = _get_bool_param(
        env,
        "document_reference_management.use_ai_analysis",
        False,
    )
    provider = str(
        _get_param(
            env,
            "document_reference_management.llm_provider",
            "ollama",
        )
    ).strip().lower()

    if not use_llm or not text.strip():
        result["ai_engine"] = "heuristic"
        return result

    try:
        if provider == "ollama":
            ai_result = analyze_with_ollama(text, env=env)
            result.update({k: v for k, v in ai_result.items() if v not in (None, "")})
            result["ai_engine"] = "ollama"
        else:
            result["ai_engine"] = "heuristic"
            result["ai_warning"] = f"Unsupported LLM provider '{provider}' in this Odoo module. Use ollama."
    except Exception as exc:
        result["ai_error"] = str(exc)
        result["ai_engine"] = "heuristic_fallback"

    return result
