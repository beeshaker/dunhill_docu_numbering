from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .branding import create_faded_logo

PLACEHOLDER = "{{REFERENCE_NO}}"


def _replace_in_paragraph(paragraph, reference_number: str) -> bool:
    replaced = False
    if PLACEHOLDER in paragraph.text:
        full_text = paragraph.text.replace(PLACEHOLDER, reference_number)
        for run in paragraph.runs:
            run.text = ""
        paragraph.add_run(full_text)
        replaced = True
    return replaced


def _replace_in_table(table, reference_number: str) -> bool:
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replaced = _replace_in_paragraph(p, reference_number) or replaced
            for nested in cell.tables:
                replaced = _replace_in_table(nested, reference_number) or replaced
    return replaced


def _add_logo_to_header(doc: Document, watermark_path: Path) -> None:
    """Best-effort Word watermark: add faint logo to each section header."""
    for section in doc.sections:
        header = section.header
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(watermark_path), width=Inches(4.8))
        except Exception:
            continue


def insert_reference_docx(
    input_path: Path,
    output_path: Path,
    reference_number: str,
    *,
    add_watermark: bool = False,
    logo_path: Path | None = None,
) -> Path:
    doc = Document(str(input_path))
    found = False

    for section in doc.sections:
        for p in section.header.paragraphs:
            found = _replace_in_paragraph(p, reference_number) or found
        for table in section.header.tables:
            found = _replace_in_table(table, reference_number) or found

    for p in doc.paragraphs:
        found = _replace_in_paragraph(p, reference_number) or found

    for table in doc.tables:
        found = _replace_in_table(table, reference_number) or found

    if not found:
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = f"Ref No.: {reference_number}"

    if add_watermark:
        watermark_path = create_faded_logo(logo_path, opacity=0.12, max_width_px=800)
        if watermark_path:
            _add_logo_to_header(doc, watermark_path)

    doc.save(str(output_path))
    return output_path
